import docx

lines = [
    "The system shall allow users to register accounts.",
    "The user enters name and password.",
    "The system validates credentials."
]

d = docx.Document()
d.add_heading("SRS", level=1)
for line in lines:
    d.add_paragraph(line)

d.save("data/srs.docx")
print("Saved data/srs.docx")